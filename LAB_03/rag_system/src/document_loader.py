"""Document loading and processing module."""

import os
from pathlib import Path
from typing import List, Dict, Tuple
import io

try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


class DocumentChunk:
    """Represents a chunk of text from a document."""

    def __init__(
        self,
        text: str,
        doc_name: str,
        chunk_id: int,
        metadata: Dict = None,
    ):
        self.text = text
        self.doc_name = doc_name
        self.chunk_id = chunk_id
        self.metadata = metadata or {}

    def __repr__(self):
        return f"DocumentChunk(doc={self.doc_name}, id={self.chunk_id}, text_len={len(self.text)})"


class DocumentLoader:
    """Loads and processes documents from various formats."""

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        """
        Initialize document loader.

        Args:
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between consecutive chunks in characters
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load_file(self, file_path: str) -> Tuple[str, str]:
        """
        Load text from a single file.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (text_content, file_name)
        """
        path = Path(file_path)
        file_name = path.name

        if path.suffix.lower() == ".pdf":
            text = self._load_pdf(file_path)
        elif path.suffix.lower() in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif path.suffix.lower() in [".docx", ".doc"]:
            text = self._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")

        return text, file_name

    def _load_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        if PdfReader is None:
            raise ImportError("PyPDF2 is required for PDF support. Install with: pip install PyPDF2")

        text = []
        with open(file_path, "rb") as f:
            pdf_reader = PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)

    def _load_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        if DocxDocument is None:
            raise ImportError(
                "python-docx is required for DOCX support. Install with: pip install python-docx"
            )

        doc = DocxDocument(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return "\n".join(text)

    def chunk_text(self, text: str, doc_name: str) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks.

        Args:
            text: Full text to chunk
            doc_name: Name of the source document

        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        chunk_id = 0

        start = 0
        while start < len(text):
            # End position for this chunk
            end = min(start + self.chunk_size, len(text))

            # Extract chunk
            chunk_text = text[start:end].strip()

            if chunk_text:  # Skip empty chunks
                chunk = DocumentChunk(
                    text=chunk_text,
                    doc_name=doc_name,
                    chunk_id=chunk_id,
                    metadata={"start_char": start, "end_char": end},
                )
                chunks.append(chunk)
                chunk_id += 1

            # Move start position for next iteration (with overlap)
            # Ensure we always move forward
            new_start = max(end - self.chunk_overlap, start + 1)
            
            # If we're at the end and haven't moved enough, break
            if new_start >= len(text):
                break
                
            start = new_start

        return chunks

    def load_and_chunk_file(self, file_path: str) -> List[DocumentChunk]:
        """
        Load a file and chunk it.

        Args:
            file_path: Path to the file

        Returns:
            List of DocumentChunk objects
        """
        text, file_name = self.load_file(file_path)
        chunks = self.chunk_text(text, file_name)
        return chunks

    def load_and_chunk_directory(self, dir_path: str) -> List[DocumentChunk]:
        """
        Load all supported files from a directory.

        Args:
            dir_path: Path to the directory

        Returns:
            List of DocumentChunk objects from all files
        """
        all_chunks = []
        path = Path(dir_path)

        supported_extensions = {".txt", ".md", ".pdf", ".docx", ".doc"}

        for file_path in path.iterdir():
            if file_path.suffix.lower() in supported_extensions:
                try:
                    chunks = self.load_and_chunk_file(str(file_path))
                    all_chunks.extend(chunks)
                    print(f"✓ Loaded {file_path.name}: {len(chunks)} chunks")
                except Exception as e:
                    print(f"✗ Error loading {file_path.name}: {e}")

        return all_chunks

    def load_from_bytes(self, file_bytes: bytes, file_name: str) -> List[DocumentChunk]:
        """
        Load document from bytes (useful for web uploads).

        Args:
            file_bytes: Document content as bytes
            file_name: Document filename

        Returns:
            List of DocumentChunk objects
        """
        suffix = Path(file_name).suffix.lower()

        if suffix == ".pdf":
            text = self._load_pdf_from_bytes(file_bytes)
        elif suffix in [".txt", ".md"]:
            text = file_bytes.decode("utf-8")
        elif suffix in [".docx", ".doc"]:
            text = self._load_docx_from_bytes(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

        chunks = self.chunk_text(text, file_name)
        return chunks

    def _load_pdf_from_bytes(self, file_bytes: bytes) -> str:
        """Extract text from PDF bytes."""
        if PdfReader is None:
            raise ImportError("PyPDF2 is required for PDF support.")

        text = []
        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n".join(text)

    def _load_docx_from_bytes(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes."""
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX support.")

        doc = DocxDocument(io.BytesIO(file_bytes))
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return "\n".join(text)
